from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOWNLOADS = Path(r"C:\Users\1\Downloads")
SOURCE = next(DOWNLOADS.glob("Разработка информационного сервис1.docx"))
OUTPUT = Path(r"C:\Users\1\EduPortal1\docs\diploma_with_project_templates_filled_v2.docx")
TEMP = Path(r"C:\Users\1\EduPortal1\scripts\_blank_templates_tmp.docx")

TOPIC = "Разработка информационного сервиса «Расписание занятий ВГТУ»"
FACULTY = "Факультет информационных технологий и компьютерной безопасности"
DEPARTMENT = "Кафедра компьютерных интеллектуальных технологий проектирования"
DIRECTION = "09.03.01 «Информатика и вычислительная техника»"
PROFILE = "«Системы автоматизированного проектирования»"
WORK_TYPE = "бакалаврская работа"
KEYWORDS = (
    "ИНФОРМАЦИОННЫЙ СЕРВИС, МОБИЛЬНОЕ ПРИЛОЖЕНИЕ, РАСПИСАНИЕ ЗАНЯТИЙ, "
    "АУДИТОРНЫЙ ФОНД, КОНСУЛЬТАЦИИ, БРОНИРОВАНИЕ АУДИТОРИЙ, "
    "EXPO, REACT NATIVE, SUPABASE"
)


def set_run_font(run, size: int = 14, bold: bool = False) -> None:
    font = run.font
    font.name = "Times New Roman"
    font.size = Pt(size)
    font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("ascii", "hAnsi", "cs"):
        rfonts.set(qn(f"w:{key}"), "Times New Roman")


def add_paragraph(doc: Document, text: str = "", *, align=WD_ALIGN_PARAGRAPH.LEFT, size: int = 14, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)


def add_blank(doc: Document, count: int = 1) -> None:
    for _ in range(count):
        add_paragraph(doc, "", size=12)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    set_run_font(r, size=12)
    r.add_break(WD_BREAK.PAGE)


def clear_section_headers_footers(section) -> None:
    for name in ("header", "footer", "first_page_header", "first_page_footer", "even_page_header", "even_page_footer"):
        part = getattr(section, name)
        part.is_linked_to_previous = False
        for paragraph in part.paragraphs:
            paragraph.text = ""


def strip_header_footer_refs(sectpr) -> None:
    for child in list(sectpr):
        if child.tag in {
            qn("w:headerReference"),
            qn("w:footerReference"),
        }:
            sectpr.remove(child)


def build_templates_doc(base: Document) -> Document:
    doc = Document()

    for section in doc.sections:
        section.page_width = base.sections[0].page_width
        section.page_height = base.sections[0].page_height
        section.left_margin = base.sections[0].left_margin
        section.right_margin = base.sections[0].right_margin
        section.top_margin = base.sections[0].top_margin
        section.bottom_margin = base.sections[0].bottom_margin
        section.header_distance = base.sections[0].header_distance
        section.footer_distance = base.sections[0].footer_distance
        section.gutter = base.sections[0].gutter
        clear_section_headers_footers(section)

    title_lines = [
        "МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        "",
        "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ",
        "ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ",
        "«ВОРОНЕЖСКИЙ ГОСУДАРСТВЕННЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ»",
        "(ФГБОУ ВО «ВГТУ», ВГТУ)",
    ]
    for line in title_lines:
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_blank(doc)
    add_paragraph(doc, FACULTY, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_paragraph(doc, DEPARTMENT, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_paragraph(doc, "Направление подготовки (специальность) " + DIRECTION, size=12)
    add_paragraph(doc, "(код, наименование)", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_paragraph(doc, "Профиль (специализация) подготовки " + PROFILE, size=12)
    add_blank(doc, 2)
    add_paragraph(doc, "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)
    add_paragraph(doc, "Бакалаврская работа", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_blank(doc, 2)
    add_paragraph(doc, "Тема " + TOPIC, size=12)
    add_blank(doc, 3)
    for line in [
        "Разработал(а)      ______________________________    _____________",
        "                   (подпись, дата)                 (инициалы, фамилия)",
        "Зав. кафедрой      ______________________________    _____________",
        "                   (подпись, дата)                 (инициалы, фамилия)",
        "Руководитель       ______________________________    _____________",
        "                   (подпись, дата)                 (инициалы, фамилия)",
        "Консультанты       ______________________________    _____________",
        "                   (подпись, дата)                 (инициалы, фамилия)",
        "                   ______________________________    _____________",
        "                   (подпись, дата)                 (инициалы, фамилия)",
    ]:
        add_paragraph(doc, line, size=12)
    add_blank(doc, 2)
    add_paragraph(doc, "Воронеж 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_page_break(doc)

    for line in title_lines:
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_blank(doc)
    add_paragraph(doc, FACULTY, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_paragraph(doc, DEPARTMENT, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_paragraph(doc, "Направление подготовки (специальность) " + DIRECTION, size=12)
    add_paragraph(doc, "Профиль (специализация) подготовки " + PROFILE, size=12)
    add_paragraph(doc, "(код, наименование)", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_paragraph(doc, "Студент группы ___________________________________________", size=12)
    add_paragraph(doc, "(индекс группы)", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_paragraph(doc, "__________________________________________________________", size=12)
    add_paragraph(doc, "(фамилия, имя, отчество)", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_blank(doc)
    add_paragraph(doc, "З А Д А Н И Е", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)
    add_blank(doc)
    add_paragraph(doc, "на выпускную квалификационную работу", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_blank(doc)
    add_paragraph(doc, "1. Тема выпускной квалификационной работы", size=12)
    add_paragraph(doc, TOPIC, size=12)
    add_paragraph(doc, "утверждена распоряжением по факультету/институту № ____ от __________ г.", size=12)
    add_blank(doc)
    add_paragraph(doc, "2. Содержание (разделы, графические работы, расчеты и проч.)", size=12)
    for line in [
        "1. Анализ предметной области и постановка задачи.",
        "2. Проектирование архитектуры, данных и пользовательских сценариев сервиса.",
        "3. Реализация мобильного приложения EduPortal на Expo и React Native.",
        "4. Реализация серверной части и хранения данных на Supabase.",
        "5. Реализация сценариев просмотра расписания, поиска преподавателя,",
        "   консультаций и бронирования аудиторий.",
        "6. Реализация алгоритма проверки конфликтов при бронировании.",
        "7. Функциональное тестирование пользовательских и административных сценариев.",
        "8. Подготовка расчетно-пояснительной записки и иллюстративных материалов.",
    ]:
        add_paragraph(doc, line, size=12)
    add_page_break(doc)

    add_paragraph(doc, "3. План выполнения выпускной квалификационной работы", size=12)
    add_paragraph(doc, "   с «      » __________________ г. по «      » __________________ г.", size=12)
    add_blank(doc)
    table = doc.add_table(rows=9, cols=4)
    table.style = "Table Grid"
    headers = ["Название элементов работы", "%", "Сроки", "Подпись руководителя, консультанта"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=11, bold=True)
    for row in table.rows[1:]:
        for cell in row.cells:
            cell.text = " "
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=11)
    add_blank(doc)
    for line in [
        "Руководитель выпускной квалификационной работы",
        "__________________________________________________________",
        "(подпись)                          (фамилия, имя, отчество)",
        "4. Выпускная квалификационная работа закончен(а)",
        "«      » __________________ г. ________________________________",
        "                              (подпись студента)",
        "5. Пояснительная записка и все материалы просмотрены",
        "Оценка руководителя __________________________________________",
        "Консультанты: ____________________   __________________________",
        "              (подпись)            (фамилия, имя, отчество)",
        "              ____________________   __________________________",
        "              (подпись)            (фамилия, имя, отчество)",
        "              ____________________   __________________________",
        "              (подпись)            (фамилия, имя, отчество)",
        "6. Допустить студента _________________________________________",
        "(фамилия, инициалы)",
        "к защите выпускной квалификационной работы в ГЭК",
        "(протокол заседания кафедры № __________ от «___»______________ г.)",
        "7. Назначить защиту на «___»__________________ г.",
        "Заведующий кафедрой _________________________________________",
        "                    (подпись)       (инициалы, фамилия)",
        "Декан факультета",
        "(Директор института) ________________________________________",
        "                     (подпись)      (инициалы, фамилия)",
    ]:
        add_paragraph(doc, line, size=11)
    add_page_break(doc)

    add_paragraph(doc, "РЕФЕРАТ", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)
    add_blank(doc, 2)
    add_paragraph(doc, "Объем ВКР ___ с., ___ рис., ___ табл., ___ источников, ___ прил.", size=12)
    add_blank(doc, 2)
    add_paragraph(doc, KEYWORDS, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_blank(doc)
    for line in [
        "Объектом исследования являются процессы получения, отображения и использования",
        "расписания занятий и данных об аудиториях в университете.",
        "Цель работы - разработка информационного сервиса «Расписание занятий ВГТУ»,",
        "обеспечивающего централизованный доступ к актуальному расписанию занятий,",
        "сведениям о преподавателях и поддержку сценария бронирования аудиторий",
        "для проведения консультаций.",
        "В процессе работы были проанализированы предметная область и существующие",
        "подходы к организации сервисов расписания, спроектированы архитектура",
        "мобильного приложения и структура данных, реализованы клиентская часть",
        "на Expo и React Native, серверное хранение данных на Supabase, а также",
        "механизм проверки конфликтов при бронировании аудитории.",
        "В результате работы создан программный прототип EduPortal, поддерживающий",
        "ролевой доступ для студента, преподавателя и администратора, просмотр",
        "расписания, поиск преподавателей, работу с консультациями, отображение",
        "состояния аудитории в планшетном режиме и административные операции.",
        "Основные технико-эксплуатационные показатели: единый мобильный интерфейс,",
        "централизованное хранение данных, ролевое разграничение доступа,",
        "предварительная проверка занятости преподавателя, группы и аудитории.",
        "Степень внедрения - разработан и протестирован работоспособный прототип.",
        "Эффективность решения заключается в сокращении числа ручных действий",
        "при получении учебной информации и снижении вероятности конфликтов",
        "при организации консультаций и использовании аудиторного фонда.",
    ]:
        add_paragraph(doc, line, size=12)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    clear_section_headers_footers(doc.sections[-1])
    for section in doc.sections:
        strip_header_footer_refs(section._sectPr)
    return doc


def merge_front(template_doc: Document, source_doc: Document, output: Path) -> None:
    body = source_doc._element.body
    original_children = list(body)
    body.clear()

    template_body = template_doc._element.body
    template_children = list(template_body)
    template_sectpr = None
    for child in template_children:
        if child.tag == qn("w:sectPr"):
            template_sectpr = deepcopy(child)
        else:
            body.append(deepcopy(child))

    for child in original_children:
        body.append(child)

    if template_sectpr is not None:
        first_p = body[0]
        ppr = first_p.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            first_p.insert(0, ppr)
        existing = ppr.find(qn("w:sectPr"))
        if existing is not None:
            ppr.remove(existing)
        ppr.append(template_sectpr)

    source_doc.save(str(output))


def main() -> None:
    copy2(SOURCE, OUTPUT)
    base_doc = Document(str(SOURCE))
    template_doc = build_templates_doc(base_doc)
    template_doc.save(str(TEMP))
    template_doc = Document(str(TEMP))
    target_doc = Document(str(OUTPUT))
    merge_front(template_doc, target_doc, OUTPUT)
    if TEMP.exists():
        TEMP.unlink()
    print(OUTPUT)


if __name__ == "__main__":
    main()
